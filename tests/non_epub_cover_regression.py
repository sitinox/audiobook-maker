#!/usr/bin/env python3
from __future__ import annotations

import importlib
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import Any, Callable, Optional

PACKAGE_PARENT = Path(__file__).resolve().parents[1]

if not (PACKAGE_PARENT / "audiobook_maker").is_dir():
    raise SystemExit(f"FAIL: Audiobook Maker package folder not found:\n{PACKAGE_PARENT}")

sys.path.insert(0, str(PACKAGE_PARENT))

def find_extractor() -> Callable[..., Any]:
    for module_name in (
        "audiobook_maker.extractors",
        "audiobook_maker.files",
        "audiobook_maker.pipeline",
        "audiobook_maker",
    ):
        try:
            module = importlib.import_module(module_name)
        except Exception:
            continue
        fn = getattr(module, "extract_source_text", None)
        if callable(fn):
            return fn
    raise SystemExit("FAIL: Could not import extract_source_text from Audiobook Maker.")

extract_source_text = find_extractor()

def extract(path: Path, title: str) -> Any:
    last_error: Optional[Exception] = None
    for args in ((path, title), (str(path), title), (path,), (str(path),)):
        try:
            return extract_source_text(*args)
        except TypeError as exc:
            last_error = exc
    raise RuntimeError(f"Could not call extract_source_text: {last_error}")

def cover_bytes(result: Any) -> Optional[bytes]:
    for name in ("cover_art", "cover"):
        value = getattr(result, name, None)

        # Audiobook Maker returns artwork as (MIME type, bytes).
        if (
            isinstance(value, tuple)
            and len(value) == 2
            and isinstance(value[1], (bytes, bytearray))
        ):
            return bytes(value[1])

        if isinstance(value, bytes):
            return value
        if isinstance(value, bytearray):
            return bytes(value)
        if isinstance(value, (str, Path)):
            p = Path(value)
            if p.exists():
                return p.read_bytes()
    return None

def detail(result: Any) -> str:
    for name in (
        "cover_detail",
        "cover_art_detail",
        "cover_source",
        "artwork_source",
        "cover_description",
    ):
        value = getattr(result, name, None)
        if value:
            return str(value).lower()

    source_type = getattr(result, "source_type", "unknown")
    has_cover = cover_bytes(result) is not None
    return f"source_type={source_type}; cover_present={has_cover}".lower()

def image_ok(data: Optional[bytes]) -> bool:
    return bool(data and len(data) > 100)

def make_cover(path: Path) -> None:
    import base64

    # Self-contained 1x1 JPEG; no Pillow or other image library required.
    jpeg_b64 = (
        "/9j/4AAQSkZJRgABAQAAAQABAAD/2wBDAP//////////////////////////////////////////////////////////////////////////////////////"
        "2wBDAf//////////////////////////////////////////////////////////////////////////////////////"
        "wAARCAABAAEDASIAAhEBAxEB/8QAFQABAQAAAAAAAAAAAAAAAAAAAAf/xAAUEAEAAAAAAAAAAAAAAAAAAAAA/"
        "9oADAMBAAIQAxAAAAF//8QAFBABAAAAAAAAAAAAAAAAAAAAAP/aAAgBAQABBQJ//8QAFBEBAAAAAAAAAAAAAAAAAAAAAP/"
        "aAAgBAwEBPwF//8QAFBEBAAAAAAAAAAAAAAAAAAAAAP/aAAgBAgEBPwF//8QAFBABAAAAAAAAAAAAAAAAAAAAAP/"
        "aAAgBAQAGPwJ//8QAFBABAAAAAAAAAAAAAAAAAAAAAP/aAAgBAQABPyF//9oADAMBAAIAAwAAABD/xAAUEQEAAAAAAAAAAAAAAAAAAAAA/"
        "9oACAEDAQE/EF//xAAUEQEAAAAAAAAAAAAAAAAAAAAA/9oACAECAQE/EF//xAAUEAEAAAAAAAAAAAAAAAAAAAAA/"
        "9oACAEBAAE/EF//2Q=="
    )
    path.write_bytes(base64.b64decode(jpeg_b64))

def make_pdf(source_txt: Path, pdf_path: Path) -> None:
    try:
        with pdf_path.open("wb") as output:
            result = subprocess.run(
                ["cupsfilter", "-m", "application/pdf", str(source_txt)],
                stdout=output,
                stderr=subprocess.PIPE,
                check=False,
            )
    except FileNotFoundError:
        raise SystemExit("FAIL: cupsfilter is unavailable on this Mac.")

    if result.returncode != 0 or not pdf_path.exists() or pdf_path.stat().st_size == 0:
        raise SystemExit(
            "FAIL: Could not create the disposable PDF: "
            + result.stderr.decode("utf-8", errors="ignore")
        )

with tempfile.TemporaryDirectory(prefix="audiobook_cover_tests_") as temp_name:
    root = Path(temp_name)
    cover = root / "generated_cover.jpg"
    make_cover(cover)
    cover_data = cover.read_bytes()

    txt_path = root / "TXT Companion Test.txt"
    txt_path.write_text("Chapter One\n\nDisposable TXT companion-cover test.", encoding="utf-8")
    txt_path.with_suffix(".jpg").write_bytes(cover_data)
    txt_result = extract(txt_path, "TXT Companion Test")

    bare_txt_path = root / "TXT No Cover Test.txt"
    bare_txt_path.write_text("Chapter One\n\nDisposable TXT with no artwork.", encoding="utf-8")
    bare_txt_result = extract(bare_txt_path, "TXT No Cover Test")

    try:
        import docx
    except ImportError as exc:
        raise SystemExit(f"FAIL: python-docx is unavailable: {exc}")

    docx_path = root / "DOCX Native Test.docx"
    document = docx.Document()
    document.add_picture(str(cover))
    document.add_paragraph("Chapter One")
    document.add_paragraph("Disposable DOCX embedded-image cover test.")
    document.save(str(docx_path))
    docx_result = extract(docx_path, "DOCX Native Test")

    docx_path.with_suffix(".jpg").write_bytes(cover_data)
    docx_override_result = extract(docx_path, "DOCX Companion Override Test")

    pdf_source = root / "PDF Native Test Source.txt"
    pdf_source.write_text(
        "PDF NATIVE COVER TEST\n\nChapter One\n\nDisposable PDF page-one test.",
        encoding="utf-8",
    )
    pdf_path = root / "PDF Native Test.pdf"
    make_pdf(pdf_source, pdf_path)
    pdf_result = extract(pdf_path, "PDF Native Test")

    pdf_path.with_suffix(".jpg").write_bytes(cover_data)
    pdf_override_result = extract(pdf_path, "PDF Companion Override Test")

    checks = [
        ("TXT companion artwork", image_ok(cover_bytes(txt_result))),
        ("TXT with no artwork", cover_bytes(bare_txt_result) is None),
        ("DOCX early embedded image", image_ok(cover_bytes(docx_result))),
        (
            "DOCX companion override",
            cover_bytes(docx_override_result) == cover_data,
        ),
        ("PDF page-one artwork", image_ok(cover_bytes(pdf_result))),
        (
            "PDF companion override",
            cover_bytes(pdf_override_result) == cover_data,
        ),
    ]

    print()
    print("NON-EPUB COVER TEST RESULTS")
    print("===========================")
    for name, passed in checks:
        print(f"{'PASS' if passed else 'FAIL'}: {name}")

    print()
    print("COVER BYTE COUNTS")
    print("-----------------")
    print("TXT companion:", len(cover_bytes(txt_result) or b""))
    print("TXT no cover:", len(cover_bytes(bare_txt_result) or b""))
    print("DOCX native:", len(cover_bytes(docx_result) or b""))
    print("DOCX override:", len(cover_bytes(docx_override_result) or b""))
    print("PDF native:", len(cover_bytes(pdf_result) or b""))
    print("PDF override:", len(cover_bytes(pdf_override_result) or b""))

    failed = [name for name, passed in checks if not passed]
    if failed:
        print()
        print("FAILED CHECKS:", ", ".join(failed))
        raise SystemExit(1)

    print()
    print("ALL NON-EPUB COVER TESTS PASSED")
    print("Disposable test files have been removed.")
